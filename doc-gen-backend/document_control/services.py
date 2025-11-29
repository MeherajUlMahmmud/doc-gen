import os
import re
from html.parser import HTMLParser
from io import BytesIO

from docx import Document
from docx.shared import RGBColor


class TemplateParser:
    PATTERN = re.compile(r'\{\{(.*?)\}\}')

    def __init__(self, file_path):
        self.file_path = file_path
        self.document = Document(file_path)

    def parse(self):
        fields = []
        seen_vars = set()
        signature_groups = {}  # Group signature fields by prefix

        def extract_from_text(text):
            matches = self.PATTERN.findall(text)
            for match in matches:
                parts = match.split('|')
                var_name = parts[0].strip()

                if var_name in seen_vars:
                    continue
                seen_vars.add(var_name)

                field_type = parts[1].strip() if len(parts) > 1 else 'text'
                label = parts[2].strip() if len(parts) > 2 else var_name.replace('_', ' ').title()
                validation = parts[3].strip() if len(parts) > 3 else ''

                field_def = {
                    'name': var_name,
                    'type': field_type,
                    'label': label,
                    'validation': validation,
                    'options': [],
                    'min_value': None,
                    'max_value': None,
                    'is_autofilled': 'autofilled' in validation.lower()
                }

                if 'options:' in validation:
                    opts = validation.split('options:')[1].split(',')
                    field_def['options'] = [opt.strip() for opt in opts]

                # Parse min/max for number fields
                if field_type == 'number':
                    if 'min:' in validation:
                        try:
                            min_part = validation.split('min:')[1].split(',')[0].split('|')[0].strip()
                            field_def['min_value'] = float(min_part)
                        except (ValueError, IndexError):
                            pass
                    if 'max:' in validation:
                        try:
                            max_part = validation.split('max:')[1].split(',')[0].split('|')[0].strip()
                            field_def['max_value'] = float(max_part)
                        except (ValueError, IndexError):
                            pass

                # Detect signature fields with pattern: prefix_N|signature|...
                if field_type == 'signature':
                    # Extract prefix and number (e.g., "requester_1" -> prefix="requester", num=1)
                    match_pattern = re.match(r'^(.+?)_(\d+)$', var_name)
                    if match_pattern:
                        prefix = match_pattern.group(1)
                        if prefix not in signature_groups:
                            signature_groups[prefix] = {
                                'prefix': prefix,
                                'base_field_name': var_name,
                                'name_field': None,
                                'section_label': label,
                                'signature_fields': [],
                                'is_required': 'required' in validation.lower()
                            }
                        signature_groups[prefix]['signature_fields'].append(field_def)
                        signature_groups[prefix]['section_label'] = label  # Use the label from signature field
                        # Update required status if any field in group is required
                        if 'required' in validation.lower():
                            signature_groups[prefix]['is_required'] = True

                # Detect associated name fields: prefix_N_name|text|...|autofilled
                if field_type == 'text' and field_def['is_autofilled']:
                    match_pattern = re.match(r'^(.+?)_(\d+)_name$', var_name)
                    if match_pattern:
                        prefix = match_pattern.group(1)
                        if prefix in signature_groups:
                            signature_groups[prefix]['name_field'] = var_name

                fields.append(field_def)

        for paragraph in self.document.paragraphs:
            extract_from_text(paragraph.text)

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        extract_from_text(paragraph.text)

        # Add signature groups metadata to return value
        result = {
            'fields': fields,
            'signature_groups': list(signature_groups.values())
        }
        return result


class HTMLToDocxConverter:
    """Convert HTML content to formatted docx runs"""

    def __init__(self, paragraph):
        self.paragraph = paragraph
        self.current_run = None
        self.formatting_stack = []

    def parse_html(self, html_content):
        """Parse HTML content and add formatted runs to paragraph"""
        if not html_content or not html_content.strip():
            return

        # Remove HTML tags if content is plain text
        if '<' not in html_content:
            self.paragraph.add_run(html_content)
            return

        # Parse HTML
        parser = HTMLFormattingParser(self.paragraph)
        parser.feed(html_content)
        parser.close()


class HTMLFormattingParser(HTMLParser):
    """Custom HTML parser to apply formatting to Word document runs"""

    def __init__(self, paragraph):
        super().__init__()
        self.paragraph = paragraph
        self.current_run = None
        self.format_stack = []
        self.current_formats = {
            'bold': False,
            'italic': False,
            'underline': False,
            'strikethrough': False,
            'color': None,
            'bg_color': None
        }

    def handle_starttag(self, tag, attrs):
        """Handle opening HTML tags"""
        attrs_dict = dict(attrs)

        if tag in ('strong', 'b'):
            self.format_stack.append(('bold', True))
            self.current_formats['bold'] = True
        elif tag in ('em', 'i'):
            self.format_stack.append(('italic', True))
            self.current_formats['italic'] = True
        elif tag == 'u':
            self.format_stack.append(('underline', True))
            self.current_formats['underline'] = True
        elif tag in ('strike', 's', 'del'):
            self.format_stack.append(('strikethrough', True))
            self.current_formats['strikethrough'] = True
        elif tag == 'span':
            # Handle inline styles
            style = attrs_dict.get('style', '')
            if 'color:' in style:
                color_match = re.search(r'color:\s*([^;]+)', style)
                if color_match:
                    color = color_match.group(1).strip()
                    self.format_stack.append(('color', color))
                    self.current_formats['color'] = color
            if 'background-color:' in style:
                bg_match = re.search(r'background-color:\s*([^;]+)', style)
                if bg_match:
                    bg_color = bg_match.group(1).strip()
                    self.format_stack.append(('bg_color', bg_color))
                    self.current_formats['bg_color'] = bg_color
        elif tag == 'br':
            # Add line break
            if self.current_run:
                self.current_run.add_break()
        elif tag == 'p':
            # Paragraph - we'll just add a line break after
            pass

    def handle_endtag(self, tag):
        """Handle closing HTML tags"""
        if tag in ('strong', 'b'):
            self._pop_format('bold')
        elif tag in ('em', 'i'):
            self._pop_format('italic')
        elif tag == 'u':
            self._pop_format('underline')
        elif tag in ('strike', 's', 'del'):
            self._pop_format('strikethrough')
        elif tag == 'span':
            # Pop color or bg_color if they were set
            if self.format_stack and self.format_stack[-1][0] in ('color', 'bg_color'):
                self._pop_format(self.format_stack[-1][0])
        elif tag == 'p':
            # Add line break after paragraph
            if self.current_run:
                self.current_run.add_break()

    def handle_data(self, data):
        """Handle text content"""
        if not data.strip():
            return

        # Create new run with current formatting
        run = self.paragraph.add_run(data)
        self.current_run = run

        # Apply formatting
        if self.current_formats['bold']:
            run.bold = True
        if self.current_formats['italic']:
            run.italic = True
        if self.current_formats['underline']:
            run.underline = True
        if self.current_formats['strikethrough']:
            run.font.strike = True

        # Apply color if set
        if self.current_formats['color']:
            rgb = self._parse_color(self.current_formats['color'])
            if rgb:
                run.font.color.rgb = RGBColor(*rgb)

    def _pop_format(self, format_type):
        """Remove format from stack"""
        if self.format_stack and self.format_stack[-1][0] == format_type:
            self.format_stack.pop()
            self.current_formats[format_type] = False

    def _parse_color(self, color_str):
        """Parse color string to RGB tuple"""
        color_str = color_str.strip().lower()

        # Handle hex colors
        if color_str.startswith('#'):
            hex_color = color_str.lstrip('#')
            if len(hex_color) == 6:
                return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))
            elif len(hex_color) == 3:
                return tuple(int(c * 2, 16) for c in hex_color)

        # Handle rgb() format
        rgb_match = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', color_str)
        if rgb_match:
            return tuple(int(x) for x in rgb_match.groups())

        # Handle common color names
        color_map = {
            'black': (0, 0, 0),
            'white': (255, 255, 255),
            'red': (255, 0, 0),
            'green': (0, 128, 0),
            'blue': (0, 0, 255),
            'yellow': (255, 255, 0),
            'cyan': (0, 255, 255),
            'magenta': (255, 0, 255),
            'gray': (128, 128, 128),
            'grey': (128, 128, 128),
        }
        return color_map.get(color_str)


class DocumentGenerator:
    def __init__(self, template_path, data, signature_path=None):
        self.template_path = template_path
        self.data = data
        self.document = Document(template_path)
        self.signature_path = signature_path

    def generate(self):
        def replace_text(paragraph):
            for key, value in self.data.items():
                if key in paragraph.text:
                    # This is a naive replacement. It might break if formatting is split across runs.
                    # For a robust solution, we'd need to iterate runs.
                    # But for now, let's try simple replacement on paragraph text.
                    # Note: modifying paragraph.text directly removes formatting.
                    # Better to iterate runs.

                    # Simple run replacement
                    for run in paragraph.runs:
                        if f'{{{{' in run.text and f'{key}' in run.text:
                            # This is also tricky because {{key}} might be split across runs.
                            pass

                    # Fallback to simple replacement if complex run handling is too much for MVP
                    # But let's try to be slightly smarter.
                    # Actually, python-docx text replacement is notoriously hard to preserve style perfectly.
                    # Let's do a simple run iteration and replace if the whole tag is in one run.
                    # If not, we might need a library or more complex logic.
                    # For this task, I will assume placeholders are likely in single runs or I will just replace paragraph text if needed,
                    # accepting some formatting loss, OR I can use a library like `docxtpl` if allowed.
                    # `docxtpl` is much better for this.
                    pass

        # Let's use a simpler approach for MVP:
        # Iterate all runs, if a run contains a full placeholder, replace it.
        # If not, we might miss it.
        # But actually, the user didn't ask for `docxtpl`.
        # Let's try to implement a basic replacement.

        for paragraph in self.document.paragraphs:
            self._replace_in_paragraph(paragraph)

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph)

        output = BytesIO()
        self.document.save(output)
        output.seek(0)
        return output

    def apply_signatures(self, signature_path):
        """Apply signature image to signature placeholders"""
        if not signature_path or not os.path.exists(signature_path):
            return

        # Find and replace signature placeholders
        for paragraph in self.document.paragraphs:
            self._replace_signature_in_paragraph(paragraph, signature_path)

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_signature_in_paragraph(paragraph, signature_path)

    def apply_multiple_signatures(self, signature_data):
        """
        Apply multiple signatures grouped by field name.
        signature_data is a dict: {field_name: [list of signature paths]}
        """
        if not signature_data:
            return

        # Find and replace signature placeholders for each field
        for paragraph in self.document.paragraphs:
            self._replace_multiple_signatures_in_paragraph(paragraph, signature_data)

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_multiple_signatures_in_paragraph(paragraph, signature_data)

    def _replace_signature_in_paragraph(self, paragraph, signature_path):
        """Replace signature placeholder with image"""
        original_text = paragraph.text
        matches = list(re.finditer(r'\{\{(.*?)\}\}', original_text))

        for match in matches:
            match_content = match.group(1)
            var_name = match_content.split('|')[0].strip()
            field_type = match_content.split('|')[1].strip() if '|' in match_content else 'text'

            if field_type == 'signature' and '{{' + match_content + '}}' in original_text:
                # Find the placeholder in runs
                placeholder = match.group(0)
                for run in paragraph.runs:
                    if placeholder in run.text:
                        # Replace placeholder with signature image
                        run.text = run.text.replace(placeholder, '')
                        # Add image after the run
                        paragraph.add_run().add_picture(signature_path, width=Inches(2))
                        break

    def _replace_multiple_signatures_in_paragraph(self, paragraph, signature_data):
        """Replace signature placeholder with multiple images based on field name"""
        original_text = paragraph.text
        matches = list(re.finditer(r'\{\{(.*?)\}\}', original_text))

        for match in matches:
            match_content = match.group(1)
            var_name = match_content.split('|')[0].strip()
            field_type = match_content.split('|')[1].strip() if '|' in match_content else 'text'

            if field_type == 'signature':
                # Check if this field name matches any in signature_data
                # Field name might be like "requester_1", we need to match by prefix
                for field_name, signature_paths in signature_data.items():
                    # Extract prefix from field_name (e.g., "requester_1" -> "requester")
                    field_prefix = re.match(r'^(.+?)_\d+$', field_name)
                    if field_prefix:
                        field_prefix = field_prefix.group(1)
                        # Check if var_name matches the prefix pattern
                        var_prefix_match = re.match(r'^(.+?)_\d+$', var_name)
                        if var_prefix_match and var_prefix_match.group(1) == field_prefix:
                            # This is a match - replace with all signatures for this field
                            placeholder = match.group(0)
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    # Replace placeholder with signature images
                                    run.text = run.text.replace(placeholder, '')
                                    # Add all signature images
                                    for sig_path in signature_paths:
                                        if os.path.exists(sig_path):
                                            paragraph.add_run().add_picture(sig_path, width=Inches(2))
                                            # Add space between signatures
                                            paragraph.add_run('  ')
                                    break
                    elif field_name == var_name:
                        # Exact match
                        placeholder = match.group(0)
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, '')
                                for sig_path in signature_paths:
                                    if os.path.exists(sig_path):
                                        paragraph.add_run().add_picture(sig_path, width=Inches(2))
                                        paragraph.add_run('  ')
                                break

    def _replace_in_paragraph(self, paragraph):
        # Get the full paragraph text to find all placeholders
        original_text = paragraph.text

        # Find all placeholders
        matches = list(re.finditer(r'\{\{(.*?)\}\}', original_text))

        if not matches:
            return

        # Collect all replacements to make
        replacements = {}
        for match in matches:
            match_content = match.group(1)
            var_name = match_content.split('|')[0].strip()
            field_type = match_content.split('|')[1].strip() if len(match_content.split('|')) > 1 else 'text'
            placeholder = match.group(0)  # Full match including {{ }}

            # Skip signature fields - they're handled separately
            if field_type == 'signature':
                continue

            if var_name in self.data:
                value = str(self.data[var_name])

                # Handle checkbox values
                if 'checkbox' in match_content.lower():
                    if value.lower() in ('true', '1', 'yes', 'on'):
                        value = 'Yes'
                    else:
                        value = 'No'

                replacements[placeholder] = value
            else:
                # If variable not in data, replace with empty string
                replacements[placeholder] = ''

        if not replacements:
            return

        # Build new paragraph text with replacements
        new_text = original_text
        for placeholder, value in replacements.items():
            new_text = new_text.replace(placeholder, value)

        if new_text == original_text:
            return

        # Preserve paragraph formatting
        paragraph_format = paragraph.paragraph_format

        # Clear the paragraph
        paragraph.clear()

        # Check if any replacement contains HTML
        has_html = any('<' in value and '>' in value for value in replacements.values())

        if has_html:
            # Use HTML parser to add formatted content
            converter = HTMLToDocxConverter(paragraph)
            converter.parse_html(new_text)
        else:
            # Simple text replacement
            if new_text:
                paragraph.add_run(new_text)

    def export_docx(self):
        """Export as DOCX format"""
        return self.generate()

    def export_doc(self):
        """Export as DOC format (legacy) - converts DOCX to DOC"""
        # Note: python-docx only works with .docx format
        # For true .doc format, we'd need additional conversion
        # For now, we'll return DOCX format but with .doc extension
        # In production, you might want to use LibreOffice or other tools
        return self.generate()

    def export_pdf(self):
        """Export as PDF format using ReportLab"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        except ImportError:
            raise Exception(
                "ReportLab is not installed. Please install it with: pip install reportlab"
            )

        # First, we need to generate the DOCX with replacements
        # But we'll work directly with the document object instead
        doc = Document(self.template_path)

        # Apply replacements to the document
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph)

        # Create PDF using ReportLab
        pdf_buffer = BytesIO()
        pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=letter,
                                    rightMargin=72, leftMargin=72,
                                    topMargin=72, bottomMargin=72)

        # Container for the 'Flowable' objects
        elements = []

        # Define styles
        styles = getSampleStyleSheet()
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=12,
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            leading=20,
            spaceAfter=12,
        )

        # Process paragraphs from DOCX
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                elements.append(Spacer(1, 0.1 * inch))
                continue

            # Check if it's a heading (simple heuristic: check style or if it's short and bold)
            is_heading = False
            if paragraph.runs:
                first_run = paragraph.runs[0]
                if first_run.bold or (len(text) < 100 and paragraph.style.name.startswith('Heading')):
                    is_heading = True

            # Clean up any remaining placeholder syntax (in case replacement missed some)
            text = re.sub(r'\{\{[^}]+\}\}', '', text)

            if text:
                # Handle basic formatting
                formatted_text = self._format_text_for_reportlab(text, paragraph)

                if is_heading:
                    elements.append(Paragraph(formatted_text, heading_style))
                else:
                    elements.append(Paragraph(formatted_text, normal_style))
                elements.append(Spacer(1, 0.1 * inch))

        # Process tables from DOCX
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = ' '.join([p.text for p in cell.paragraphs]).strip()
                    # Clean up any remaining placeholders
                    cell_text = re.sub(r'\{\{[^}]+\}\}', '', cell_text)
                    row_data.append(cell_text if cell_text else '')
                table_data.append(row_data)

            if table_data:
                # Create ReportLab table
                pdf_table = Table(table_data)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                ]))
                elements.append(Spacer(1, 0.2 * inch))
                elements.append(pdf_table)
                elements.append(Spacer(1, 0.2 * inch))

        # Build PDF
        pdf_doc.build(elements)
        pdf_buffer.seek(0)
        return pdf_buffer

    def _format_text_for_reportlab(self, text, paragraph):
        """Convert paragraph text with basic formatting to ReportLab format"""
        # ReportLab uses XML-like tags for formatting
        # Try to preserve bold and italic from runs
        if paragraph.runs:
            parts = []
            for run in paragraph.runs:
                run_text = run.text
                if not run_text:
                    continue
                # Escape XML special characters first
                run_text = run_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                # Apply formatting tags
                if run.bold and run.italic:
                    run_text = f"<b><i>{run_text}</i></b>"
                elif run.bold:
                    run_text = f"<b>{run_text}</b>"
                elif run.italic:
                    run_text = f"<i>{run_text}</i>"
                parts.append(run_text)
            if parts:
                return ''.join(parts)

        # Fallback: just escape the text
        formatted = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        return formatted
