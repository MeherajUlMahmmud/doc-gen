import logging

from django.db.models import Q
from django.http import FileResponse
from django.utils import timezone
from django.contrib.auth import get_user_model
from rest_framework import status
from rest_framework.response import Response

from common.custom_view import (
    CustomListAPIView, CustomWOPListAPIView, CustomRetrieveAPIView, CustomCreateAPIView, CustomDestroyAPIView
)
from .models import GeneratedDocumentModel, DocumentSignatoryModel, TemplateModel
from .serializers.document_signatory import DocumentSignatoryModelSerializer
from .serializers.generated_document import GeneratedDocumentModelSerializer
from .serializers.template import TemplateModelSerializer
from user_control.serializers.user import UserModelSerializer

logger = logging.getLogger(__name__)
User = get_user_model()


class PendingSignaturesAPIView(CustomWOPListAPIView):
    """
    Get all documents pending the authenticated user's signature
    Returns JSON data for React frontend
    """
    serializer_class = DocumentSignatoryModelSerializer.List

    def get_queryset(self):
        """Get all DocumentSignatory records where user is a pending signatory"""
        return DocumentSignatoryModel.objects.filter(
            signatory=self.request.user,
            status='pending'
        ).select_related('document', 'document__template', 'document__user').order_by('-document__created_at')

    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset()
        serializer = self.get_serializer(queryset, many=True)
        return Response({
            'status': 'success',
            'message': 'Pending signatures retrieved successfully',
            'data': serializer.data
        }, status=status.HTTP_200_OK)


class ApproveDocumentAPIView(CustomCreateAPIView):
    """
    Approve and sign a document - handles individual signatory signing
    Expects: { "pin": "1234", "totp_code": "123456" (optional) }
    """
    serializer_class = GeneratedDocumentModelSerializer.Approve
    lookup_field = 'pk'
    lookup_url_kwarg = 'pk'

    def get_queryset(self):
        """Get all documents (for lookup)"""
        return GeneratedDocumentModel.objects.all()

    def post(self, request, pk, *args, **kwargs):
        # Get the document
        self.kwargs['pk'] = pk
        document = self.get_object()

        # Find the signatory record for this user
        try:
            signatory = DocumentSignatoryModel.objects.get(
                document=document,
                signatory=request.user,
                status='pending'
            )
        except DocumentSignatoryModel.DoesNotExist:
            return Response({
                'status': 'error',
                'message': 'You are not a pending signatory for this document'
            }, status=status.HTTP_403_FORBIDDEN)

        # Validate input using serializer
        serializer = self.get_serializer(data=request.data)
        if not serializer.is_valid():
            return Response({
                'status': 'error',
                'message': 'Validation failed',
                'errors': serializer.errors
            }, status=status.HTTP_400_BAD_REQUEST)

        validated_data = serializer.validated_data
        pin = validated_data.get('pin')
        totp_code = validated_data.get('totp_code')

        # Verify PIN
        if not request.user.check_pin(pin):
            return Response({
                'status': 'error',
                'message': 'Invalid PIN'
            }, status=status.HTTP_400_BAD_REQUEST)

        # Verify 2FA if enabled and code provided
        if request.user.two_factor_enabled and totp_code:
            if not request.user.verify_totp(totp_code):
                return Response({
                    'status': 'error',
                    'message': 'Invalid 2FA code'
                }, status=status.HTTP_400_BAD_REQUEST)

        # Check if user has a signature file
        if not request.user.signature_file:
            return Response({
                'status': 'error',
                'message': 'Please upload a signature in your profile first'
            }, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Update signatory status
            signatory.status = 'approved'
            signatory.signed_at = timezone.now()
            signatory.save()

            logger.info(f"Document {document.id} signed by {request.user.email}")

            # Check if all signatories have signed
            pending_count = DocumentSignatoryModel.objects.filter(
                document=document,
                status='pending'
            ).count()

            if pending_count == 0:
                # All signatories have signed - update document status
                document.status = 'approved'
                document.save()
                logger.info(f"Document {document.id} fully approved - all signatories have signed")

            # Serialize the updated document for output
            # Use Detail serializer for full document information
            output_serializer = GeneratedDocumentModelSerializer.Detail(
                document,
                context={'request': request}
            )

            return Response({
                'status': 'success',
                'message': 'Document signed successfully',
                'data': output_serializer.data
            }, status=status.HTTP_200_OK)

        except Exception as e:
            logger.error(f"Error approving document {pk}: {str(e)}")
            return Response({
                'status': 'error',
                'message': f'Failed to approve document: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class TemplatesListAPIView(CustomListAPIView):
    """
    Get all active templates with pagination and filtering
    Query params:
    - search: Search in title and description
    - is_active: Filter by active status (true/false)
    - page: Page number (default: 1)
    - page_size: Items per page (default: 12)
    """
    serializer_class = TemplateModelSerializer.List

    def get_queryset(self):
        """Build query with search and active status filters"""
        queryset = TemplateModel.objects.all().order_by('-created_at')

        # Get query parameters
        search = self.request.GET.get('search', '').strip()
        is_active = self.request.GET.get('is_active', 'true').lower() == 'true'

        # Apply active status filter
        queryset = queryset.filter(is_active=is_active)

        # Apply search filter
        if search:
            queryset = queryset.filter(
                Q(title__icontains=search) | Q(description__icontains=search)
            )

        return queryset

    def list(self, request, *args, **kwargs):
        response = super().list(request, *args, **kwargs)
        # Override to match the expected response format
        page = int(request.GET.get('page', 1))
        page_size = int(request.GET.get('page_size', 12))

        return Response({
            'status': 'success',
            'message': 'Templates retrieved successfully',
            'data': response.data.get('data', []),
            'meta': {
                'total_records': response.data.get('total_records', 0),
                'total_pages': response.data.get('total_pages', 0),
                'current_page': page,
                'page_size': page_size,
            }
        }, status=status.HTTP_200_OK)


class TemplateDetailAPIView(CustomRetrieveAPIView):
    """
    Get a single template by ID
    """
    serializer_class = TemplateModelSerializer.Detail
    lookup_field = 'pk'
    lookup_url_kwarg = 'pk'

    def get_queryset(self):
        """Get active templates only"""
        return TemplateModel.objects.filter(is_active=True)

    def retrieve(self, request, *args, **kwargs):
        instance = self.get_object()
        serializer = self.get_serializer(instance)
        data = serializer.data.copy()
        
        # Parse template file to extract fields
        try:
            from .services import TemplateParser
            import os
            from django.conf import settings
            
            template_path = instance.file.path
            if os.path.exists(template_path):
                parser = TemplateParser(template_path)
                parsed_data = parser.parse()
                
                # Add fields and signature groups to response
                data['fields'] = parsed_data.get('fields', [])
                data['signature_groups'] = parsed_data.get('signature_groups', [])
                
                # Mark required fields
                for field in data['fields']:
                    field['required'] = 'required' in field.get('validation', '').lower()
        except Exception as e:
            logger.warning(f"Failed to parse template fields for template {instance.id}: {str(e)}")
            data['fields'] = []
            data['signature_groups'] = []
        
        return Response({
            'status': 'success',
            'message': 'Template retrieved successfully',
            'data': data
        }, status=status.HTTP_200_OK)


class TemplateVersionsAPIView(CustomRetrieveAPIView):
    """
    Get all versions of a template by ID
    """
    serializer_class = TemplateModelSerializer.Detail
    lookup_field = 'pk'
    lookup_url_kwarg = 'pk'

    def get_queryset(self):
        """Get all templates (including inactive) for version listing"""
        return TemplateModel.objects.all()

    def retrieve(self, request, *args, **kwargs):
        instance = self.get_object()
        # Get all versions using the model method
        all_versions = instance.get_all_versions()
        serializer = self.get_serializer(all_versions, many=True)
        return Response({
            'status': 'success',
            'message': 'Template versions retrieved successfully',
            'data': serializer.data
        }, status=status.HTTP_200_OK)


class TemplateUploadAPIView(CustomCreateAPIView):
    """
    Upload a new template
    Requires: multipart/form-data with fields:
    - title: Template title
    - description: Template description (optional)
    - file: Template file (.docx)
    - upload_type: 'new' or 'version'
    - parent_template_id: Parent template ID (required when upload_type is 'version')
    """
    serializer_class = TemplateModelSerializer.Create

    def post(self, request, *args, **kwargs):
        # Check if user is admin
        if not request.user.is_staff:
            return Response({
                'status': 'error',
                'message': 'Only administrators can upload templates'
            }, status=status.HTTP_403_FORBIDDEN)

        # Validate using serializer
        serializer = self.get_serializer(data=request.data)
        if not serializer.is_valid():
            return Response({
                'status': 'error',
                'message': 'Validation failed',
                'errors': serializer.errors
            }, status=status.HTTP_400_BAD_REQUEST)

        try:
            validated_data = serializer.validated_data
            upload_type = validated_data.get('upload_type')
            parent_template_id = validated_data.get('parent_template_id')
            title = validated_data.get('title')
            description = validated_data.get('description', '')
            file = validated_data.get('file')

            if upload_type == 'version':
                # Create a new version of existing template
                parent_template = TemplateModel.objects.get(id=parent_template_id)
                latest_version = parent_template.get_latest_version()

                template = TemplateModel.objects.create(
                    title=title,
                    description=description,
                    file=file,
                    parent=parent_template,
                    version=latest_version.version + 1
                )
                message = f'New version (v{template.version}) uploaded successfully'
            else:
                # Create a new template
                template = TemplateModel.objects.create(
                    title=title,
                    description=description,
                    file=file,
                    version=1,
                    parent=None
                )
                message = 'New template uploaded successfully'

            output_serializer = TemplateModelSerializer.Detail(template)

            return Response({
                'status': 'success',
                'message': message,
                'data': output_serializer.data
            }, status=status.HTTP_201_CREATED)

        except TemplateModel.DoesNotExist:
            return Response({
                'status': 'error',
                'message': 'Parent template not found'
            }, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            logger.error(f"Error uploading template: {str(e)}")
            return Response({
                'status': 'error',
                'message': f'Failed to upload template: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class DocumentCreateAPIView(CustomCreateAPIView):
    """
    Create a new document from template with signatory assignments
    Expects JSON: {
        "name": "Document title",
        "template": "template_id",
        "input_data": {...},
        "export_format": "docx",  // optional
        "signatories": [  // optional
            {
                "user_id": "uuid",
                "signature_field_name": "requester_1",
                "signature_order": 0
            },
            ...
        ]
    }
    """
    serializer_class = GeneratedDocumentModelSerializer.Create

    def post(self, request, *args, **kwargs):
        # Validate document data using serializer
        serializer = self.get_serializer(data=request.data)
        if not serializer.is_valid():
            return Response({
                'status': 'error',
                'message': 'Validation failed',
                'errors': serializer.errors
            }, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Create document
            document = serializer.save(user=request.user)

            # Process signatory assignments if provided
            signatories_data = request.data.get('signatories', [])
            if signatories_data:
                signatory_objects = []
                for sig_data in signatories_data:
                    user_id = sig_data.get('user_id')
                    signature_field_name = sig_data.get('signature_field_name')
                    signature_order = sig_data.get('signature_order', 0)

                    # Validate user exists
                    try:
                        signatory_user = User.objects.get(id=user_id)
                    except User.DoesNotExist:
                        continue

                    # Create signatory assignment
                    signatory_obj = DocumentSignatoryModel(
                        document=document,
                        signatory=signatory_user,
                        signature_field_name=signature_field_name,
                        signature_order=signature_order,
                        status='pending'
                    )
                    signatory_objects.append(signatory_obj)

                # Bulk create signatory assignments
                if signatory_objects:
                    DocumentSignatoryModel.objects.bulk_create(signatory_objects)

                    # Update document counts
                    document.signatories_count = len(signatory_objects)
                    document.status = 'pending_approval'
                    document.save()

                    logger.info(f"Document {document.id} created with {len(signatory_objects)} signatories")

            # Return document details
            output_serializer = GeneratedDocumentModelSerializer.Detail(
                document,
                context={'request': request}
            )

            return Response({
                'status': 'success',
                'message': 'Document created successfully',
                'data': output_serializer.data
            }, status=status.HTTP_201_CREATED)

        except Exception as e:
            logger.error(f"Error creating document: {str(e)}")
            return Response({
                'status': 'error',
                'message': f'Failed to create document: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class DocumentsListAPIView(CustomListAPIView):
    """
    Get all documents for the authenticated user with pagination and filtering
    Query params:
    - status: Filter by document status (draft, generated, pending_approval, approved)
    - search: Search in document name and template title
    - page: Page number (default: 1)
    - page_size: Items per page (default: 12)
    """
    serializer_class = GeneratedDocumentModelSerializer.List

    def get_queryset(self):
        """Build query with status and search filters"""
        queryset = GeneratedDocumentModel.objects.filter(
            user=self.request.user
        ).select_related('template', 'user').order_by('-created_at')

        # Get query parameters
        search = self.request.GET.get('search', '').strip()
        status_filter = self.request.GET.get('status', '').strip()

        # Apply status filter
        if status_filter:
            queryset = queryset.filter(status=status_filter)

        # Apply search filter
        if search:
            queryset = queryset.filter(
                Q(name__icontains=search) | Q(template__title__icontains=search)
            )

        return queryset

    def list(self, request, *args, **kwargs):
        response = super().list(request, *args, **kwargs)
        # Override to match the expected response format
        page = int(request.GET.get('page', 1))
        page_size = int(request.GET.get('page_size', 12))

        return Response({
            'status': 'success',
            'message': 'Documents retrieved successfully',
            'data': response.data.get('data', []),
            'meta': {
                'total_records': response.data.get('total_records', 0),
                'total_pages': response.data.get('total_pages', 0),
                'current_page': page,
                'page_size': page_size,
            }
        }, status=status.HTTP_200_OK)


class DocumentDetailAPIView(CustomRetrieveAPIView):
    """
    Get a single document by ID
    """
    serializer_class = GeneratedDocumentModelSerializer.Detail
    lookup_field = 'pk'
    lookup_url_kwarg = 'pk'

    def get_queryset(self):
        """Get documents for the authenticated user only"""
        return GeneratedDocumentModel.objects.filter(user=self.request.user)

    def retrieve(self, request, *args, **kwargs):
        instance = self.get_object()
        serializer = self.get_serializer(instance)
        return Response({
            'status': 'success',
            'message': 'Document retrieved successfully',
            'data': serializer.data
        }, status=status.HTTP_200_OK)


class DocumentDeleteAPIView(CustomDestroyAPIView):
    """
    Delete a document
    """
    lookup_field = 'pk'
    lookup_url_kwarg = 'pk'

    def get_queryset(self):
        """Get documents for the authenticated user only"""
        return GeneratedDocumentModel.objects.filter(user=self.request.user)

    def destroy(self, request, *args, **kwargs):
        instance = self.get_object()
        self.perform_destroy(instance)
        return Response({
            'status': 'success',
            'message': 'Document deleted successfully'
        }, status=status.HTTP_200_OK)


class TemplateDownloadAPIView(CustomRetrieveAPIView):
    """
    Download a template file
    """
    queryset = TemplateModel.objects.filter(is_active=True)
    lookup_field = 'pk'
    lookup_url_kwarg = 'pk'

    def retrieve(self, request, *args, **kwargs):
        template = self.get_object()

        if not template.file:
            return Response({
                'status': 'error',
                'message': 'Template file not found'
            }, status=status.HTTP_404_NOT_FOUND)

        # Return the file as a download
        response = FileResponse(
            template.file.open('rb'),
            content_type='application/octet-stream'
        )
        response['Content-Disposition'] = f'attachment; filename="{template.file.name}"'

        return response


class TemplatePreviewAPIView(CustomCreateAPIView):
    """
    Generate a preview of the template with filled data
    POST /api/v1/templates/<uuid:pk>/preview/
    Body: { "fields": { "field_name": "value", ... } }
    Returns: HTML preview of the document
    """
    queryset = TemplateModel.objects.filter(is_active=True)
    lookup_field = 'pk'
    lookup_url_kwarg = 'pk'

    def get_object(self):
        """Get the template object from the URL parameter"""
        pk = self.kwargs.get(self.lookup_url_kwarg)
        return self.queryset.get(pk=pk)

    def post(self, request, *args, **kwargs):
        template = self.get_object()

        if not template.file:
            return Response({
                'status': 'error',
                'message': 'Template file not found'
            }, status=status.HTTP_404_NOT_FOUND)

        try:
            from .services import DocumentGenerator
            from docx import Document
            import tempfile
            import os

            # Get the fields data from request
            fields_data = request.data.get('fields', {})

            # Generate document with the provided data
            generator = DocumentGenerator(template.file.path, fields_data)
            doc_bytes = generator.generate()

            # Save to a temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                tmp_file.write(doc_bytes.getvalue())
                tmp_path = tmp_file.name

            try:
                # Convert to HTML for preview
                doc = Document(tmp_path)
                html_content = self._convert_to_html(doc)

                return Response({
                    'status': 'success',
                    'message': 'Preview generated successfully',
                    'data': {
                        'html': html_content,
                        'fields': fields_data
                    }
                }, status=status.HTTP_200_OK)
            finally:
                # Clean up temporary file
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)

        except Exception as e:
            logger.error(f"Error generating preview for template {template.id}: {str(e)}")
            return Response({
                'status': 'error',
                'message': f'Failed to generate preview: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def _convert_to_html(self, doc):
        """Convert docx document to HTML"""
        html_parts = ['<div class="document-preview" style="font-family: Arial, sans-serif; line-height: 1.6; padding: 20px; max-width: 800px; margin: 0 auto;">']

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Determine paragraph style
                style_attrs = []
                if paragraph.style.name.startswith('Heading'):
                    html_parts.append(f'<h3 style="margin-top: 1em; margin-bottom: 0.5em; font-weight: bold;">{self._escape_html(paragraph.text)}</h3>')
                else:
                    # Check for bold/italic runs
                    if any(run.bold for run in paragraph.runs):
                        html_parts.append(f'<p style="margin: 0.5em 0;"><strong>{self._escape_html(paragraph.text)}</strong></p>')
                    elif any(run.italic for run in paragraph.runs):
                        html_parts.append(f'<p style="margin: 0.5em 0;"><em>{self._escape_html(paragraph.text)}</em></p>')
                    else:
                        html_parts.append(f'<p style="margin: 0.5em 0;">{self._escape_html(paragraph.text)}</p>')

        # Handle tables
        for table in doc.tables:
            html_parts.append('<table style="width: 100%; border-collapse: collapse; margin: 1em 0;">')
            for row in table.rows:
                html_parts.append('<tr>')
                for cell in row.cells:
                    cell_text = ' '.join([p.text for p in cell.paragraphs if p.text.strip()])
                    html_parts.append(f'<td style="border: 1px solid #ddd; padding: 8px;">{self._escape_html(cell_text)}</td>')
                html_parts.append('</tr>')
            html_parts.append('</table>')

        html_parts.append('</div>')
        return ''.join(html_parts)

    def _escape_html(self, text):
        """Escape HTML special characters"""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;'))


class DocumentDownloadAPIView(CustomRetrieveAPIView):
    """
    Download a document file
    """
    lookup_field = 'pk'
    lookup_url_kwarg = 'pk'

    def get_queryset(self):
        """Get documents for the authenticated user only"""
        return GeneratedDocumentModel.objects.filter(user=self.request.user)

    def retrieve(self, request, *args, **kwargs):
        document = self.get_object()

        if not document.generated_file:
            return Response({
                'status': 'error',
                'message': 'Document file not found'
            }, status=status.HTTP_404_NOT_FOUND)

        # Return the file as a download
        response = FileResponse(
            document.generated_file.open('rb'),
            content_type='application/octet-stream'
        )
        response['Content-Disposition'] = f'attachment; filename="{document.generated_file.name}"'

        return response


class UsersListAPIView(CustomListAPIView):
    """
    Get all active users for signatory selection
    Query params:
    - search: Search in name and email
    - division: Filter by division
    - designation: Filter by designation
    """
    serializer_class = UserModelSerializer.Lite

    def get_queryset(self):
        """Build query with search and filters"""
        queryset = User.objects.filter(
            is_active=True,
            is_verified=True
        ).order_by('first_name', 'last_name')

        # Get query parameters
        search = self.request.GET.get('search', '').strip()
        division = self.request.GET.get('division', '').strip()
        designation = self.request.GET.get('designation', '').strip()

        # Apply division filter
        if division:
            queryset = queryset.filter(division__icontains=division)

        # Apply designation filter
        if designation:
            queryset = queryset.filter(designation__icontains=designation)

        # Apply search filter
        if search:
            queryset = queryset.filter(
                Q(first_name__icontains=search) |
                Q(last_name__icontains=search) |
                Q(email__icontains=search)
            )

        return queryset

    def list(self, request, *args, **kwargs):
        """Return list of users without pagination"""
        queryset = self.get_queryset()
        serializer = self.get_serializer(queryset, many=True)

        return Response({
            'status': 'success',
            'message': 'Users retrieved successfully',
            'data': serializer.data
        }, status=status.HTTP_200_OK)
